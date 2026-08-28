from __future__ import annotations

from io import BytesIO

import pandas as pd

from app.ingestion.extraction_contract import (
    ArtifactExtractionResult,
    ExtractedDataset,
    ExtractedEvidence,
)

# See powerpoint_parser.py's module docstring comment: python-docx (and its
# lxml dependency) is optional at runtime -- its unavailability must never
# prevent the app from starting or unrelated parsers from working.
_IMPORT_ERROR: str | None = None
try:
    from docx import Document
except ImportError as exc:  # pragma: no cover -- environment dependent
    _IMPORT_ERROR = str(exc)


class WordParser:
    code = "docx"
    version = "1.0"

    def supports(self, mime_type: str, extension: str) -> bool:
        return extension.lower() == ".docx" or mime_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    def is_available(self) -> bool:
        return _IMPORT_ERROR is None

    def extract(self, raw_bytes: bytes, filename: str) -> ArtifactExtractionResult:
        if _IMPORT_ERROR is not None:
            return ArtifactExtractionResult(
                parser_code=self.code,
                parser_version=self.version,
                status="unavailable",
                warnings=[
                    "DOCX parsing dependency (python-docx/lxml) is unavailable in this "
                    f"environment: {_IMPORT_ERROR}"
                ],
                extraction_metadata={"reason": "dependency_unavailable"},
            )
        try:
            document = Document(BytesIO(raw_bytes))
        except Exception as exc:  # noqa: BLE001
            return ArtifactExtractionResult(
                parser_code=self.code,
                parser_version=self.version,
                status="failed",
                warnings=[f"DOCX parse failed: {exc}"],
            )
        paragraphs = "\n".join(p.text for p in document.paragraphs if p.text.strip())
        evidence = []
        if paragraphs:
            evidence.append(ExtractedEvidence(evidence_type="text_block", content=paragraphs))
        core = document.core_properties
        metadata = {
            "title": core.title,
            "author": core.author,
            "created": core.created.isoformat() if core.created else None,
        }
        evidence.append(
            ExtractedEvidence(evidence_type="metadata", content=str(metadata), lineage=metadata)
        )
        datasets = []
        for table_index, table in enumerate(document.tables):
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            if len(rows) < 2:
                continue
            dataframe = pd.DataFrame(rows[1:], columns=rows[0])
            datasets.append(
                ExtractedDataset(
                    label=f"{filename}:table_{table_index + 1}",
                    dataframe=dataframe,
                    lineage={"table_index": table_index + 1},
                )
            )
        return ArtifactExtractionResult(
            parser_code=self.code,
            parser_version=self.version,
            status="extracted",
            evidence_objects=evidence,
            datasets=datasets,
            extraction_metadata=metadata,
        )
